import io
from abc import ABC, abstractmethod

import docx
from fastapi import UploadFile
from pypdf import PdfReader
from docx2python import docx2python

from ...models.CVModels import PartialCVData


class AbstractCVParser(ABC):

    def parse(self, firstname: str, lastname: str, file: UploadFile) -> PartialCVData:
        cv = self._parse(firstname, lastname, file)
        cv.firstname = firstname
        cv.lastname = lastname
        return cv


    def _parse(self, firstname:str, lastname: str, file: UploadFile) -> PartialCVData:
        if file.filename.endswith('.pdf'):
            return self._parse_cv(self.convert_pdf_to_str(file), firstname, lastname)
        elif file.filename.endswith('.docx'):
            return self._parse_cv(self.convert_docx_to_str(file), firstname, lastname)
        elif file.filename.endswith('.txt'):
            return self._parse_cv(self.convert_txt_to_str(file), firstname, lastname)
        else:
            raise Exception('Unsupported file format')

    def convert_pdf_to_str(self, file: UploadFile) -> str:
        text = ''
        reader = PdfReader(file.file)
        pageCount = len(reader.pages)
        pageNum = 0
        while pageNum < pageCount:
            page = reader.pages[pageNum]
            text += page.extract_text()
            pageNum += 1
        return text

    def convert_docx_to_str(self, file: UploadFile) -> str:
        ioBytes = io.BytesIO(file.file.read())
        # Old version
        # doc = docx.Document(ioBytes)
        # fullText = []
        # for para in doc.paragraphs:
        #     fullText.append(para.text)
        # return '\n'.join(fullText)
        text = ""
        doc = docx.Document(ioBytes)
        with docx2python(ioBytes) as docx_content:
            text = docx_content.text
        return text
        

    def convert_txt_to_str(self, file: UploadFile) -> str:
        return file.file.read().decode('utf-8')

    @abstractmethod
    def _parse_cv(self, text: str, firstname: str, lastname: str) -> PartialCVData:
        pass
